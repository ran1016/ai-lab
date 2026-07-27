"""简历解析模块：从 PDF / DOCX 提取文本，并结构化映射到站点板块。

解析策略：
1. 若配置了 LLM_API_KEY，优先调用 OpenAI 兼容接口做高精度结构化（中文简历友好）。
2. 否则回退到本地启发式规则（正则提取邮箱/电话/姓名 + 按区块关键字分段）。

返回结构统一为：
{
  "profile": {"name", "headline", "bio", "email", "phone", "location"},
  "sections": {
      "experience": [{"title","subtitle","description","tags"}],
      "skill":      [{"title","subtitle","description","tags"}],
      "project":    [{"title","subtitle","description","tags","external_url"}],
  },
  "engine": "llm" | "heuristic"
}
"""
import json
import os
import re

# ----------------------- 文本提取 -----------------------


def extract_text(path):
    """根据扩展名提取简历纯文本。"""
    ext = path.rsplit(".", 1)[-1].lower() if "." in path else ""
    if ext == "pdf":
        return _extract_pdf(path)
    if ext == "docx":
        return _extract_docx(path)
    if ext == "doc":
        raise ValueError("暂不支持旧版 .doc 格式，请在 Word 中另存为 .docx 或导出 PDF")
    raise ValueError("不支持的文件类型（仅支持 PDF / DOCX）")


def _extract_pdf(path):
    try:
        import pdfplumber
    except ImportError:
        raise ValueError("未安装 pdfplumber，无法解析 PDF（请执行 pip install pdfplumber）")
    chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            if text:
                chunks.append(text)
    result = "\n".join(chunks).strip()
    if not result:
        raise ValueError("未能从 PDF 中提取到文本（可能是扫描件图片，请使用可复制文字的 PDF）")
    return result


def _extract_docx(path):
    try:
        import docx
    except ImportError:
        raise ValueError("未安装 python-docx，无法解析 DOCX（请执行 pip install python-docx）")
    document = docx.Document(path)
    paras = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))
    return "\n".join(paras).strip()


# ----------------------- LLM 结构化 -----------------------


SYSTEM_PROMPT = (
    "你是一个专业的简历解析助手。请从用户提供的简历文本中提取结构化信息，"
    "并严格只输出一个 JSON 对象（不要使用 markdown 代码块，不要任何额外解释）。"
    "JSON 结构如下：\n"
    "{\n"
    '  "profile": {\n'
    '    "name": "姓名（未知留空字符串）",\n'
    '    "headline": "求职意向 / 目标职位 / 一句话头衔",\n'
    '    "bio": "个人总结或自我评价（2-4 句话，浓缩为一段）",\n'
    '    "email": "邮箱地址",\n'
    '    "phone": "电话号码",\n'
    '    "location": "所在地或城市"\n'
    "  },\n"
    '  "sections": {\n'
    '    "experience": [{"title":"职位/经历名称","subtitle":"公司或机构及时间段","description":"要点描述","tags":"相关技术,逗号分隔"}],\n'
    '    "skill": [{"title":"单个技能名称(若简历用逗号/顿号列举多种技能,请拆分为多个独立对象,不要合并成一条)","subtitle":"熟练度或所属分类(如 熟练/了解/精通,或 编程语言/框架)","description":"","tags":""}],\n'
    '    "project": [{"title":"项目名称","subtitle":"担任角色或时间","description":"项目描述与成果","tags":"技术栈,逗号分隔","external_url":""}]\n'
    "  }\n"
    "}\n"
    "只填写能确定的字段，无法确定则留空字符串或空数组。"
    "experience / project 条目按时间倒序（最近的在前）。"
)


def parse_with_llm(text, api_key, base_url, model):
    """调用 OpenAI 兼容接口解析简历文本，返回结构化 dict。"""
    import urllib.request

    url = base_url.rstrip("/") + "/chat/completions"
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": "以下是简历文本：\n\n" + text},
        ],
        "temperature": 0.1,
        "response_format": {"type": "json_object"},
    }
    data = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(
        url,
        data=data,
        headers={
            "Authorization": "Bearer " + api_key,
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=90) as resp:
            body = json.loads(resp.read().decode("utf-8"))
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"LLM 接口调用失败：{exc}")

    content = body["choices"][0]["message"]["content"]
    return _coerce(_extract_json(content))


def _extract_json(text):
    text = text.strip()
    # 去除可能的 ```json ... ``` 包裹
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1:
        raise ValueError("LLM 未返回有效 JSON")
    return json.loads(text[start : end + 1])


# ----------------------- 启发式规则解析（兜底） -----------------------


SECTION_KEYWORDS = {
    "experience": ["工作经历", "实习经历", "工作/实习", "实践经历", "项目与经历"],
    "project": ["项目经历", "项目经验", "主要项目", "项目"],
    "skill": ["专业技能", "技能特长", "技能", "技术栈"],
    "education": ["教育背景", "教育经历", "教育"],
}


def parse_heuristic(text):
    """无 LLM 时的本地解析：尽力提取基础信息并按区块分段。"""
    profile = {}
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    m = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    if m:
        profile["email"] = m.group(0)
    m = re.search(r"(?:(?:\+?86[-\s]?)?)(1[3-9]\d{9})", text)
    if m:
        profile["phone"] = m.group(1)
    m = re.search(r"(?:姓名|名字)\s*[:：]\s*([^\s,，]{1,12})", text)
    if m:
        profile["name"] = m.group(1)
    m = re.search(r"(?:求职意向|目标职位|期望职位|应聘职位)\s*[:：]\s*([^\n]{1,40})", text)
    if m:
        profile["headline"] = m.group(1)
    m = re.search(r"(?:所在地|城市|现居|地点)\s*[:：]\s*([^\n]{1,20})", text)
    if m:
        profile["location"] = m.group(1)
    for label in ("自我评价", "个人总结", "自我描述", "个人简介"):
        m = re.search(label + r"[:：]?\s*\n(.*?)(?:\n[A-Z一-龥]{2,6}[:：]|$)", text, re.S)
        if m:
            profile["bio"] = re.sub(r"\s+", " ", m.group(1)).strip()[:400]
            break

    sections = _split_sections(text)
    return {
        "profile": profile,
        "sections": sections,
        "engine": "heuristic",
    }


def _split_sections(text):
    """按区块关键字把简历正文切分成结构化条目。"""
    result = {"experience": [], "skill": [], "project": []}
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]

    current = None
    buffer = []
    # 把"教育"并入 experience（作为经历），便于展示时间线
    mapped = {}
    for cat, keys in SECTION_KEYWORDS.items():
        for k in keys:
            mapped[k] = cat

    def flush(cat, buf):
        if cat is None or not buf:
            return
        items = _lines_to_items(buf, cat)
        target = "experience" if cat == "education" else cat
        result.setdefault(target, []).extend(items)

    for line in lines:
        matched_cat = None
        for k, cat in mapped.items():
            if k in line and len(line) <= len(k) + 8:
                matched_cat = cat
                break
        if matched_cat:
            flush(current, buffer)
            current = matched_cat
            buffer = []
        else:
            buffer.append(line)
    flush(current, buffer)

    for cat in list(result):
        result[cat] = [i for i in result[cat] if i.get("title")]
    return result


# ----------------------- 技能解析（重点优化） -----------------------


_SKILL_CAT_WORDS = (
    "编程语言", "语言", "框架", "工具", "数据库", "平台", "前端", "后端",
    "中间件", "云", "开发", "技术栈", "专业技能", "技能",
)
_SKILL_LEVELS = {"精通", "熟练", "了解", "掌握", "熟悉", "良好", "一般", "熟练使用"}


def _split_skills(s):
    """把被逗号/顿号/分号等分隔的技能字符串拆成多个技能名。"""
    parts = re.split(r"[，,、;；/\\|]", s)
    out = []
    for p in parts:
        p = p.strip().strip("。；;").strip()
        if p and len(p) <= 30:
            out.append(p)
    return out


def _skill_lines_to_items(lines):
    """把技能区块文本解析为多个独立技能条目，支持：
    - 分类式：编程语言：Python、Java
    - 熟练度式：Python - 熟练 / Python：熟练
    - 逗号列举：Python、Java、SQL
    - 列表逐项
    """
    items = []
    seen = set()

    def add(title, subtitle=""):
        title = title.strip()
        if not title or len(title) > 40:
            return
        if title in seen:
            # 已存在：熟练度信息优先补入 subtitle
            if subtitle:
                for it in items:
                    if it["title"] == title:
                        cur = it["subtitle"]
                        new_level = subtitle in _SKILL_LEVELS
                        cur_level = cur in _SKILL_LEVELS
                        if not cur or (new_level and not cur_level):
                            it["subtitle"] = subtitle[:120]
            return
        seen.add(title)
        items.append(
            {"title": title[:120], "subtitle": subtitle[:120], "description": "", "tags": title}
        )

    for raw in lines:
        line = re.sub(r"^[\s]*([•\-*·▸▹◦]|[0-9]+[.、])\s*", "", raw).strip()
        if not line:
            continue
        # 1) 分类：值（如「编程语言：Python、Java」）
        m = re.match(r"^([一-龥A-Za-z]{1,12})[：:]\s*(.+)$", line)
        if m and (
            m.group(1) in _SKILL_CAT_WORDS
            or m.group(1).endswith(("语言", "框架", "工具", "技能"))
        ):
            for s in _split_skills(m.group(2)):
                add(s, m.group(1))
            continue
        # 2) 技能 - 熟练度（如「Python - 熟练」）
        if " - " in line:
            a, b = line.split(" - ", 1)
            for s in _split_skills(a):
                add(s, b.strip())
            continue
        # 3) 技能：熟练度（如「Python：熟练」）
        if re.search(r"[：:]\s*(熟练|了解|精通|掌握|熟悉|良好|一般|熟练使用)", line):
            parts = re.split(r"[：:]", line, 1)
            for s in _split_skills(parts[0]):
                add(s, parts[1].strip())
            continue
        # 4) 逗号 / 顿号列举的多个技能（如「Python、Java、SQL」）
        skills = _split_skills(line)
        if len(skills) > 1:
            for s in skills:
                add(s)
            continue
        # 5) 单行单技能
        add(line)
    return items


def _lines_to_items(lines, cat=None):
    """将一段文本行拆为条目列表；cat 为 'skill' 时走技能专用解析。"""
    if cat == "skill":
        return _skill_lines_to_items(lines)

    items = []
    blocks = []
    cur = []
    for ln in lines:
        if re.match(r"^[\s]*([•\-*·▸▹◦]|[0-9]+[.、])\s*", ln) or (
            re.match(r"^[0-9]{4}", ln) and ln.endswith(("：", ":", "年"))
        ):
            if cur:
                blocks.append(cur)
            cur = [ln]
        else:
            cur.append(ln)
    if cur:
        blocks.append(cur)

    if not blocks:
        blocks = [lines]

    for block in blocks:
        if not block:
            continue
        first = re.sub(r"^[\s]*([•\-*·▸▹◦]|[0-9]+[.、])\s*", "", block[0]).strip()
        if not first:
            continue
        # 尝试拆出 title / subtitle（用 - 或 ：分割）
        title, subtitle = first, ""
        if " - " in first:
            title, subtitle = first.split(" - ", 1)
        elif "：" in first or ":" in first:
            parts = re.split(r"[:：]", first, 1)
            title, subtitle = parts[0].strip(), parts[1].strip()
        desc = " ".join(block[1:]).strip()
        desc = re.sub(r"^[\s]*([•\-*·▸▹◦]|[0-9]+[.、])\s*", "", desc)
        tags = ",".join(re.findall(r"(?:Python|Java|SQL|C\+\+|Go|Rust|Vue|React|PyTorch|TensorFlow|MySQL|Redis|Docker|Linux|Excel|Tableau|Hadoop|Spark|LLM|RAG|Agent|FastAPI|Flask)", " ".join(block)))
        items.append(
            {
                "title": title[:120],
                "subtitle": subtitle[:120],
                "description": desc[:600],
                "tags": tags,
            }
        )
    return items


# ----------------------- 统一入口 -----------------------


def _coerce(data):
    """把 LLM 返回的数据规整成统一结构。"""
    profile = data.get("profile") or {}
    sections = data.get("sections") or {}
    out = {
        "profile": {k: (profile.get(k) or "") for k in ("name", "headline", "bio", "email", "phone", "location")},
        "sections": {
            "experience": [],
            "skill": [],
            "project": [],
        },
        "engine": data.get("engine", "llm"),
    }
    for cat in ("experience", "skill", "project"):
        for it in sections.get(cat, []) or []:
            if not it or not it.get("title"):
                continue
            item = {
                "title": str(it.get("title") or "")[:200],
                "subtitle": str(it.get("subtitle") or "")[:200],
                "description": str(it.get("description") or "")[:2000],
                "tags": str(it.get("tags") or "")[:300],
            }
            if cat == "project":
                item["external_url"] = str(it.get("external_url") or "")[:500]
            out["sections"][cat].append(item)
    return out


def parse_resume(text, api_key=None, base_url=None, model=None):
    """解析入口：有 key 走 LLM，否则走规则。"""
    if api_key:
        try:
            return parse_with_llm(text, api_key, base_url or "https://api.openai.com/v1", model or "gpt-4o-mini")
        except Exception as exc:  # noqa: BLE001
            raise  # 让调用方决定是否回退
    return parse_heuristic(text)
